"""
DOCX 简历解析模块
使用 python-docx 提取 Word 文档文本
"""

from docx import Document


def extract_text_from_docx(docx_file) -> str:
    """
    从 DOCX 文件提取文本
    
    参数:
        docx_file: Streamlit 上传的文件对象或文件路径
    
    返回:
        提取的文本内容
    """
    try:
        # 读取 DOCX 文件
        if hasattr(docx_file, 'read'):
            # Streamlit 上传的文件对象
            doc = Document(docx_file)
        else:
            # 文件路径
            doc = Document(docx_file)
        
        text_parts = []
        
        # 遍历每个段落提取文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # 遍历表格提取文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        # 合并所有文本
        full_text = "\n".join(text_parts)
        
        # 简单清理：移除多余空行
        lines = full_text.split('\n')
        cleaned_lines = []
        for line in lines:
            cleaned_line = line.strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        return "\n".join(cleaned_lines)
    
    except Exception as e:
        print(f"DOCX 解析失败: {e}")
        return f"❌ DOCX 解析失败: {str(e)}"
